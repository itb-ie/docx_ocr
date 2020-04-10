import pdf2image
import pytesseract

import docx
from docx.shared import Pt


# This function will create a docx document based on the text we are passing. Each text is a page
def create_doc(doc, text):
    para_texts = text.split("\n\n")
    for p_text in para_texts:
        para = doc.add_paragraph()
        p_text = p_text.replace("\n", " ")
        run = para.add_run(p_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


# We need to tell pytesseract where we have installed the tesseract application!
pytesseract.pytesseract.tesseract_cmd = "C:/Program Files/Tesseract-OCR/tesseract"
pdf_file = "scansmpl"

# create the empty doc
doc = docx.Document()

# first step is to convert the pdf to image. The result is a list of images, each image is a page from the PDf file
images = pdf2image.convert_from_path(pdf_file+".pdf")

# go page by page (image in images)
for image in images:
    # convert one page to text using the tesseract OCR:
    text = pytesseract.image_to_string(image, lang="eng")

    # now make it a page in our docx
    create_doc(doc, text)

    #add a page break, but do not add if we are on the last image :)
    if image != images[-1]:
        doc.add_page_break()


# save the doc
doc.save(pdf_file+".docx")